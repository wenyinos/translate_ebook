#!/usr/bin/env python3
"""
电子书翻译脚本 - 使用 OpenAI 兼容端点
支持 .docx 和 .epub 格式
"""

import os
import sys
import json
import time
import argparse
from pathlib import Path
from typing import List, Tuple, Optional

# 文档处理
from docx import Document
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup

# OpenAI API
import openai


# ==================== 配置 ====================
# 通过环境变量或命令行参数配置
DEFAULT_CONFIG = {
    "api_key": os.environ.get("OPENAI_API_KEY", ""),
    "base_url": os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1"),
    "model": os.environ.get("OPENAI_MODEL", "gpt-4o-mini"),
    "batch_size": int(os.environ.get("TRANSLATE_BATCH_SIZE", "50")),
    "max_retries": int(os.environ.get("TRANSLATE_MAX_RETRIES", "3")),
    "retry_delay": float(os.environ.get("TRANSLATE_RETRY_DELAY", "1.0")),
    "output_dir": os.environ.get("TRANSLATE_OUTPUT_DIR", str(Path.home() / "translated_books")),
}


# ==================== 翻译核心 ====================
def translate_text(client: openai.OpenAI, text: str, model: str, max_retries: int = 3, retry_delay: float = 1.0) -> str:
    """使用 OpenAI API 翻译文本"""
    if not text.strip():
        return text

    # 跳过纯数字、标点、页码等
    if text.strip().isdigit() or len(text.strip()) <= 2:
        return text

    prompt = f"""请将以下英文文本翻译为简体中文。
要求：
1. 保持原文格式和排版
2. 技术术语准确（如 SoC、CPU、GPU、API 等保留原文或使用通用译法）
3. 章节编号保持不变
4. 只输出翻译结果，不要添加解释

原文：
{text}"""

    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是一个专业的技术文档翻译员，擅长将英文技术文档翻译为简体中文。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=4096
            )
            result = response.choices[0].message.content
            return result.strip() if result else text
        except Exception as e:
            if attempt < max_retries - 1:
                print(f"  翻译失败，重试 {attempt + 1}/{max_retries}: {e}")
                time.sleep(retry_delay)
            else:
                print(f"  翻译失败，跳过: {e}")
                return text  # 失败时返回原文


def translate_batch(client: openai.OpenAI, texts: List[str], model: str, batch_size: int = 50, max_workers: int = 4) -> List[str]:
    """批量翻译文本（支持并行）"""
    from concurrent.futures import ThreadPoolExecutor, as_completed
    import threading

    results = [None] * len(texts)
    total = len(texts)
    completed = 0
    lock = threading.Lock()
    start_time = time.time()

    def format_time(seconds):
        """格式化时间显示"""
        if seconds < 60:
            return f"{seconds:.0f}秒"
        elif seconds < 3600:
            return f"{seconds/60:.1f}分钟"
        else:
            return f"{seconds/3600:.1f}小时"

    def translate_item(idx_text):
        idx, text = idx_text
        translated = translate_text(client, text, model)
        with lock:
            nonlocal completed
            completed += 1
            elapsed = time.time() - start_time
            speed = completed / elapsed if elapsed > 0 else 0
            remaining = (total - completed) / speed if speed > 0 else 0
            percent = (completed / total) * 100
            
            # 每 10 个或完成时显示进度
            if completed % 10 == 0 or completed == total:
                print(f"  进度: {completed}/{total} ({percent:.1f}%) | "
                      f"速度: {speed:.2f}个/秒 | "
                      f"已用: {format_time(elapsed)} | "
                      f"剩余: {format_time(remaining)}")
        return idx, translated

    # 使用线程池并行翻译
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(translate_item, (i, text)) for i, text in enumerate(texts)]
        for future in as_completed(futures):
            idx, translated = future.result()
            results[idx] = translated

    return results


# ==================== DOCX 处理 ====================
def translate_docx(input_path: str, output_path: str, client: openai.OpenAI, model: str, batch_size: int = 50):
    """翻译 DOCX 文件"""
    print(f"\n处理 DOCX: {input_path}")

    # 读取源文档
    src_doc = Document(input_path)
    total_paragraphs = len(src_doc.paragraphs)
    print(f"  总段落数: {total_paragraphs}")

    # 提取需要翻译的段落
    paragraphs_to_translate = []
    for i, para in enumerate(src_doc.paragraphs):
        if para.text.strip():
            paragraphs_to_translate.append((i, para.text))

    total_to_translate = len(paragraphs_to_translate)
    print(f"  非空段落数: {total_to_translate}")

    start_time = time.time()

    # 批量翻译
    texts = [p[1] for p in paragraphs_to_translate]
    translated_texts = translate_batch(client, texts, model, batch_size)

    # 创建新文档
    dst_doc = Document()

    # 复制所有段落（包括空段落）
    translation_idx = 0
    for i, para in enumerate(src_doc.paragraphs):
        new_para = dst_doc.add_paragraph()

        if para.text.strip():
            # 有内容的段落，使用翻译后的文本
            new_para.text = translated_texts[translation_idx]
            translation_idx += 1
        else:
            # 空段落保持为空
            new_para.text = ""

        # 复制样式
        if para.style:
            new_para.style = para.style

        # 复制格式
        if para.alignment:
            new_para.alignment = para.alignment

    # 保存
    dst_doc.save(output_path)
    total_time = time.time() - start_time
    print(f"  翻译完成: {output_path}")
    print(f"  总耗时: {total_time:.1f}秒")


# ==================== EPUB 处理 ====================
def html_to_text(html_content: str) -> str:
    """从 HTML 提取纯文本"""
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup.get_text(separator='\n', strip=True)


def text_to_html(text: str, original_html: str) -> str:
    """将翻译后的文本转换回 HTML（保持原结构）"""
    soup = BeautifulSoup(original_html, 'html.parser')

    # 找到所有文本节点并替换
    for string in soup.find_all(string=True):
        if string.strip():
            # 这里简化处理，实际可能需要更复杂的逻辑
            pass

    # 简单方案：用翻译后的文本替换 body 内容
    body = soup.find('body')
    if body:
        # 保持原有的 HTML 结构，只替换文本
        for child in body.children:
            if child.name:
                child.string = text
                break

    return str(soup)


def translate_epub(input_path: str, output_path: str, client: openai.OpenAI, model: str, batch_size: int = 50):
    """翻译 EPUB 文件"""
    print(f"\n处理 EPUB: {input_path}")

    # 读取源 epub
    book = epub.read_epub(input_path)

    # 获取所有文档项
    items = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
    total_items = len(items)
    print(f"  文档项数量: {total_items}")

    start_time = time.time()

    def format_time(seconds):
        """格式化时间显示"""
        if seconds < 60:
            return f"{seconds:.0f}秒"
        elif seconds < 3600:
            return f"{seconds/60:.1f}分钟"
        else:
            return f"{seconds/3600:.1f}小时"

    # 翻译每个文档项
    for idx, item in enumerate(items):
        elapsed = time.time() - start_time
        speed = (idx + 1) / elapsed if elapsed > 0 else 0
        remaining = (total_items - idx - 1) / speed if speed > 0 else 0
        percent = ((idx + 1) / total_items) * 100

        print(f"  处理文档项 {idx + 1}/{total_items} ({percent:.1f}%) | "
              f"速度: {speed:.2f}个/秒 | "
              f"已用: {format_time(elapsed)} | "
              f"剩余: {format_time(remaining)}")

        # 提取文本
        html_content = item.get_content().decode('utf-8')
        text = html_to_text(html_content)

        if not text.strip():
            continue

        # 翻译文本
        translated_text = translate_text(client, text, model)

        # 将翻译后的文本放回 HTML
        # 简单方案：创建一个包含翻译文本的新 HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        body = soup.find('body')
        if body:
            # 清空 body 内容，添加翻译后的文本
            for child in list(body.children):
                child.decompose()

            # 按段落分割翻译后的文本
            paragraphs = translated_text.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p_tag = soup.new_tag('p')
                    p_tag.string = para_text
                    body.append(p_tag)

        item.set_content(str(soup).encode('utf-8'))

    # 保存
    epub.write_epub(output_path, book)
    total_time = time.time() - start_time
    print(f"  翻译完成: {output_path}")
    print(f"  总耗时: {format_time(total_time)}")


# ==================== 主程序 ====================
def main():
    parser = argparse.ArgumentParser(description='电子书翻译脚本 - 使用 OpenAI 兼容端点')
    parser.add_argument('input', help='输入文件路径')
    parser.add_argument('-o', '--output', help='输出文件路径')
    parser.add_argument('--output-dir', help='输出目录（默认: ~/translated_books）')
    parser.add_argument('--api-key', help='OpenAI API Key')
    parser.add_argument('--base-url', help='OpenAI API Base URL')
    parser.add_argument('--model', help='模型名称')
    parser.add_argument('--batch-size', type=int, help='批量翻译大小')
    parser.add_argument('--max-retries', type=int, help='最大重试次数')
    parser.add_argument('--retry-delay', type=float, help='重试延迟（秒）')

    args = parser.parse_args()

    # 合并配置
    config = DEFAULT_CONFIG.copy()
    if args.api_key:
        config["api_key"] = args.api_key
    if args.base_url:
        config["base_url"] = args.base_url
    if args.model:
        config["model"] = args.model
    if args.batch_size:
        config["batch_size"] = args.batch_size
    if args.max_retries:
        config["max_retries"] = args.max_retries
    if args.retry_delay:
        config["retry_delay"] = args.retry_delay
    if args.output_dir:
        config["output_dir"] = args.output_dir

    # 检查 API Key
    if not config["api_key"]:
        print("错误: 请设置 OPENAI_API_KEY 环境变量或使用 --api-key 参数")
        sys.exit(1)

    # 检查输入文件
    input_path = args.input
    if not os.path.exists(input_path):
        print(f"错误: 输入文件不存在: {input_path}")
        sys.exit(1)

    # 确定输出路径
    if args.output:
        output_path = args.output
    else:
        input_name = Path(input_path).name
        output_path = str(Path(config["output_dir"]) / input_name)

    # 创建输出目录
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # 初始化 OpenAI 客户端
    client = openai.OpenAI(
        api_key=config["api_key"],
        base_url=config["base_url"]
    )

    # 根据文件类型调用相应的翻译函数
    file_ext = Path(input_path).suffix.lower()

    if file_ext == '.docx':
        translate_docx(
            input_path, output_path, client,
            config["model"], config["batch_size"]
        )
    elif file_ext == '.epub':
        translate_epub(
            input_path, output_path, client,
            config["model"], config["batch_size"]
        )
    else:
        print(f"错误: 不支持的文件格式: {file_ext}")
        sys.exit(1)

    print(f"\n翻译完成!")


if __name__ == "__main__":
    main()
