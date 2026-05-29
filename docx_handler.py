"""DOCX 文件处理 - 保留段落结构、内联格式和表格"""

from typing import Optional, List, Tuple
from docx import Document
from docx.table import Table
from translator import translate_batch, TokenStats, KeyManager, TranslationCache, DEFAULT_TARGET_LANG


def copy_run_format(src_run, dst_run):
    """复制单个 run 的格式（粗体、斜体、下划线等）"""
    if src_run.bold:
        dst_run.bold = True
    if src_run.italic:
        dst_run.italic = True
    if src_run.underline:
        dst_run.underline = True
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.name:
        dst_run.font.name = src_run.font.name


def copy_paragraph_format(src_para, dst_para):
    """复制段落格式（对齐、缩进等）"""
    if src_para.alignment:
        dst_para.alignment = src_para.alignment
    pf = src_para.paragraph_format
    if pf.first_line_indent:
        dst_para.paragraph_format.first_line_indent = pf.first_line_indent
    if pf.left_indent:
        dst_para.paragraph_format.left_indent = pf.left_indent
    if pf.right_indent:
        dst_para.paragraph_format.right_indent = pf.right_indent


def translate_docx(input_path: str, output_path: str, client, model: str,
                   batch_size: int = 50, resume: bool = False,
                   token_stats: Optional[TokenStats] = None,
                   target_lang: str = DEFAULT_TARGET_LANG,
                   max_tokens: int = 128000,
                   key_manager: Optional[KeyManager] = None,
                   cache: Optional[TranslationCache] = None):
    """翻译 DOCX 文件，保留样式、内联格式和表格"""
    print(f"\nProcessing DOCX: {input_path}")

    src_doc = Document(input_path)

    # 收集所有需要翻译的文本（段落 + 表格）
    all_texts: List[Tuple[str, str]] = []  # (type, text)

    # 段落文本
    for i, para in enumerate(src_doc.paragraphs):
        if para.text.strip():
            all_texts.append(('paragraph', para.text))

    # 表格文本
    for table_idx, table in enumerate(src_doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    all_texts.append(('table', cell.text))

    total_paragraphs = len(src_doc.paragraphs)
    total_tables = len(src_doc.tables)
    total_to_translate = len([t for t in all_texts if t[1].strip()])
    print(f"  Stats: {total_paragraphs} paragraphs, {total_tables} tables, {total_to_translate} items to translate")

    # 批量翻译
    texts = [t[1] for t in all_texts]
    translated_texts = translate_batch(client, texts, model, batch_size,
                                       max_tokens=max_tokens,
                                       output_path=output_path, resume=resume,
                                       token_stats=token_stats,
                                       target_lang=target_lang,
                                       key_manager=key_manager, cache=cache)

    # 创建新文档并填充内容
    dst_doc = Document()

    # 复制段落
    translation_idx = 0
    for i, para in enumerate(src_doc.paragraphs):
        new_para = dst_doc.add_paragraph()

        if para.text.strip():
            translated_text = translated_texts[translation_idx]
            translation_idx += 1

            if len(para.runs) > 1:
                translated_runs = split_text_to_runs(translated_text, len(para.runs))
                for j, (src_run, run_text) in enumerate(zip(para.runs, translated_runs)):
                    dst_run = new_para.add_run(run_text)
                    copy_run_format(src_run, dst_run)
            else:
                dst_run = new_para.add_run(translated_text)
                if para.runs:
                    copy_run_format(para.runs[0], dst_run)
        else:
            new_para.text = ""

        if para.style:
            new_para.style = para.style
        copy_paragraph_format(para, new_para)

    # 复制表格
    for table_idx, src_table in enumerate(src_doc.tables):
        # 创建新表格
        rows = len(src_table.rows)
        cols = len(src_table.columns)
        dst_table = dst_doc.add_table(rows=rows, cols=cols)

        # 复制表格样式
        if src_table.style:
            dst_table.style = src_table.style

        # 填充表格内容
        for row_idx, row in enumerate(src_table.rows):
            for cell_idx, cell in enumerate(row.cells):
                dst_cell = dst_table.cell(row_idx, cell_idx)

                if cell.text.strip():
                    translated_text = translated_texts[translation_idx]
                    translation_idx += 1

                    # 复制单元格内段落格式
                    if cell.paragraphs:
                        src_para = cell.paragraphs[0]
                        dst_para = dst_cell.paragraphs[0]

                        if len(src_para.runs) > 1:
                            translated_runs = split_text_to_runs(translated_text, len(src_para.runs))
                            for j, (src_run, run_text) in enumerate(zip(src_para.runs, translated_runs)):
                                dst_run = dst_para.add_run(run_text)
                                copy_run_format(src_run, dst_run)
                            # 清除默认空段落
                            if dst_para.runs and dst_para.text == '':
                                dst_para.text = ''
                        else:
                            dst_para.text = translated_text
                            if src_para.runs:
                                copy_run_format(src_para.runs[0], dst_para.runs[0] if dst_para.runs else dst_para.add_run(''))

    dst_doc.save(output_path)
    print(f"  Translation complete: {output_path}")


def split_text_to_runs(text: str, num_runs: int) -> list:
    """将翻译文本按 run 数量分割"""
    if num_runs <= 1:
        return [text]

    # 按空格或标点分割
    import re
    words = re.findall(r'\S+|\s+', text)

    if len(words) <= num_runs:
        # 词数少于 run 数，均匀分配
        result = [''] * num_runs
        for i, word in enumerate(words):
            result[i % num_runs] += word
        return result
    else:
        # 词数多于 run 数，按比例分配
        words_per_run = len(words) // num_runs
        remainder = len(words) % num_runs
        result = []
        idx = 0
        for i in range(num_runs):
            count = words_per_run + (1 if i < remainder else 0)
            result.append(''.join(words[idx:idx + count]))
            idx += count
        return result
